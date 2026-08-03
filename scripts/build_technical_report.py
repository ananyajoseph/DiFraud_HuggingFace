"""Create the technical report from the successfully executed notebook outputs."""

from __future__ import annotations

import base64
import json
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
NOTEBOOK = ROOT / "difraud_pu_semisupervised_analysis.ipynb"
OUTPUT = ROOT / "DiFrauD_Technical_Report.docx"

NAVY = "1F4D78"
BLUE = "2E74B5"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
MUTED = "666666"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths))); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd"); tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def font(run, size=11, bold=False, color="000000", italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar"); fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar"); fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_table(doc, headers, rows, widths, numeric_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, label in enumerate(headers):
        set_cell_shading(hdr.cells[i], LIGHT)
        p = hdr.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(label), 9, True, NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(str(value)), 9)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(8)
    font(p.add_run(text), 9, False, MUTED, True)


def add_callout(doc, label, text, color=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1); table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [9360]); set_cell_shading(table.cell(0, 0), color)
    p = table.cell(0, 0).paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    font(p.add_run(label + "  "), 10, True, NAVY)
    font(p.add_run(text), 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.5)
    p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    font(p.add_run(text), 10.5)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), 10.5, True)
        font(p.add_run(text[len(bold_lead):]), 10.5)
    else:
        font(p.add_run(text), 10.5)
    return p


def notebook_figures():
    nb = json.loads(NOTEBOOK.read_text(encoding="utf-8"))
    figures = []
    for cell in nb["cells"]:
        for out in cell.get("outputs", []):
            payload = out.get("data", {}).get("image/png")
            if payload:
                figures.append(BytesIO(base64.b64decode(payload)))
    return figures


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8); section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1); section.right_margin = Inches(1)
section.header_distance = Inches(.492); section.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor(0, 0, 0)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, NAVY, 8, 4)):
    style = styles[name]; style.font.name = "Calibri"; style.font.size = Pt(size)
    style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
font(header.add_run("DIFRAUD TECHNICAL REPORT"), 8.5, True, MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(footer.add_run("DiFrauD PU and Semi-Supervised Study  |  "), 8.5, False, MUTED)
add_page_field(footer)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(3)
font(p.add_run("TECHNICAL REPORT"), 10, True, BLUE)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
font(p.add_run("Positive-Unlabeled and Mean Teacher Learning for Cross-Domain Deception Detection"), 23, True, NAVY)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
font(p.add_run("A leakage-resistant quick-mode study on the seven-domain DiFrauD benchmark"), 13, False, MUTED)

meta = [
    ("Primary evidence", "Successfully executed notebook; 14/14 code cells, zero execution errors"),
    ("Dataset revision", "aaaf94b336c563a14806bb4f3f58727bed9ed8d4"),
    ("Execution scope", "CPU quick mode; 10,500 stratified records; one seed (42)"),
    ("Evaluation", "Official splits plus reduced LODO: job scams, product reviews, phishing"),
    ("Date", "August 2, 2026"),
]
add_table(doc, ["Report field", "Value"], meta, [2100, 7260])
add_callout(doc, "Evidence boundary", "All numerical claims in this report come from the stored outputs of the successfully executed quick-mode notebook. Full three-seed, all-domain LODO, and transformer fine-tuning configurations were not executed and are not presented as results.")

add_heading(doc, "Executive summary", 1)
add_body(doc, "This study evaluated whether deception signals transfer across seven heterogeneous text domains when only a small share of deceptive training examples is labeled. It compared a valid non-negative positive-unlabeled estimator (nnPU), an invalid unlabeled-as-negative shortcut, a Mean Teacher consistency extension, supervised baselines, and cross-domain holdouts. The central empirical result is cautionary: in-domain text models learned useful signal, but domain provenance was highly predictable and cross-domain performance was weak or unstable.")
add_body(doc, "The fully supervised TF-IDF logistic-regression upper bound achieved F1 0.712 and PR-AUC 0.775 on the official test splits. At 10% revealed positives, nnPU achieved F1 0.628 and PR-AUC 0.603, narrowly exceeding the naive comparator on those metrics. The combined PU-Mean Teacher model reached F1 0.617 and PR-AUC 0.577, so consistency learning did not improve nnPU at that operating point in this one-seed run. At 20%, nnPU reached F1 0.653 and the combined model reached F1 0.650.")
add_body(doc, "A text-only domain classifier achieved 93.1% accuracy. Metadata-only and domain-prior baselines were also predictive, demonstrating that the benchmark contains strong source and composition cues. The near-duplicate audit found 296 pairs above cosine similarity 0.92 in the quick sample, including 212 cross-split pairs. These observations materially limit claims that the models detect deception itself rather than source-specific artifacts.")

add_heading(doc, "1. Research objective and questions", 1)
add_body(doc, "The project develops and evaluates a domain-independent textual deception detector under positive-label scarcity. The target setting assumes that a small subset of deceptive examples is labeled while the remaining training pool mixes hidden positives and genuine negatives. The study asks:")
for item in (
    "How much performance remains when 1%, 5%, 10%, or 20% of deceptive training examples are revealed?",
    "Does correct nnPU learning outperform the common but invalid practice of treating all unlabeled examples as negative?",
    "Does confidence-filtered teacher-student consistency improve a PU detector?",
    "Do text models learn deception, or domain identity and constituent-dataset artifacts?",
    "How well does the detector generalize to a completely held-out domain?",
): add_bullet(doc, item)

add_heading(doc, "2. Dataset and execution scope", 1)
add_body(doc, "DiFrauD contains English binary text-classification data from phishing, fake news, political statements, product reviews, job scams, SMS, and Twitter rumours. The official repository exposes 21 JSONL files: train, validation, and test for each domain. Direct JSONL loading avoided execution of the repository's remote loading script. The full snapshot contained 95,854 records with no missing text, empty text, missing labels, or invalid binary labels.")
domain_rows = [
    ("Fake news", "16,364", "2,046", "2,046", "43.2%"),
    ("Job scams", "11,436", "1,429", "1,430", "4.2%"),
    ("Phishing", "12,217", "1,527", "1,528", "39.8%"),
    ("Political statements", "9,997", "1,250", "1,250", "64.3%"),
    ("Product reviews", "16,776", "2,097", "2,098", "50.0%"),
    ("SMS", "5,259", "657", "658", "19.4%"),
    ("Twitter rumours", "4,631", "579", "579", "34.0%"),
]
add_table(doc, ["Domain", "Train", "Validation", "Test", "Train prevalence"], domain_rows,
          [2600, 1500, 1700, 1500, 2060], numeric_cols=(1, 2, 3, 4))
add_body(doc, "Quick mode retained 500 stratified examples per domain and official split, producing 10,500 analyzed records. This preserves official split membership but is a representative computational reduction, not a full-dataset experiment. Frozen all-MiniLM-L6-v2 sentence embeddings and small PyTorch multilayer perceptrons enabled CPU execution.")

figures = notebook_figures()
if figures:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(figures[0], width=Inches(6.35))
    shape._inline.docPr.set("descr", "Bar charts showing quick-mode sample composition and deceptive prevalence across the seven DiFrauD domains.")
    add_caption(doc, "Figure 1. Quick-mode composition and deceptive prevalence by domain. Stored notebook output.")

add_heading(doc, "3. Dataset-quality audit", 1)
add_heading(doc, "3.1 Composition, integrity, and imbalance", 2)
add_body(doc, "Class prevalence varied from approximately 4.2% for job scams to 64.3% for political statements. A model can therefore gain substantial apparent performance by learning domain identity. The audit measured character, word, sentence, URL, digit, punctuation, capitalization, non-ASCII, repeated-punctuation, HTML-residue, and replacement-character features. The full acquisition check found zero missing or empty texts, but integrity and provenance signals remained domain-dependent.")
add_heading(doc, "3.2 Duplicate and leakage risk", 2)
add_body(doc, "No exact normalized duplicate groups appeared in the 10,500-row quick sample. Character 3-5 gram TF-IDF nearest-neighbor screening identified 296 near-duplicate pairs at cosine similarity at least 0.92; 212 crossed official split boundaries. Representative pairs included structurally similar account-verification emails and related news or conversational text. Because near-duplicate auditing was performed on the quick sample, these counts are lower bounds for the full repository.")
add_callout(doc, "Leakage implication", "Cross-split near duplicates can inflate in-domain performance even when the official split labels are preserved. A publication-grade full run should de-duplicate or cluster before splitting and report sensitivity to cluster-level exclusion.", "FFF4E5")
add_heading(doc, "3.3 Shortcut learning and domain separability", 2)
shortcut_rows = [
    ("Metadata-only logistic regression", "F1", "0.576"),
    ("Metadata-only logistic regression", "ROC-AUC", "0.655"),
    ("Domain-prior baseline", "F1", "0.613"),
    ("Domain-prior baseline", "ROC-AUC", "0.721"),
    ("Text domain classifier", "Accuracy", "0.931"),
]
add_table(doc, ["Diagnostic model", "Metric", "Result"], shortcut_rows, [5000, 2200, 2160], numeric_cols=(2,))
add_body(doc, "The 93.1% domain-classifier accuracy is the clearest shortcut warning: domain provenance is much easier to recover than a domain-independent deception concept. Domain-prior and metadata-only baselines further show that composition and shallow formatting carry label information. Out-of-fold label-disagreement analysis surfaced ambiguous job-scam positives and Twitter-rumour or SMS negatives, but these are suspected label issues rather than confirmed annotation errors.")
if len(figures) > 1:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(figures[1], width=Inches(6.35))
    shape._inline.docPr.set("descr", "Two TF-IDF SVD scatter plots colored first by domain and then by recorded deception label.")
    add_caption(doc, "Figure 2. TF-IDF/SVD representation colored by domain and recorded label. Stored notebook output.")

add_heading(doc, "4. Experimental design", 1)
add_heading(doc, "4.1 Leakage-resistant evaluation", 2)
add_body(doc, "In-domain experiments trained only on official training rows, selected thresholds and early-stopping states on validation rows, and used test labels only for final scoring. LODO experiments removed the held-out domain from training, prior estimation, early stopping, calibration, and threshold selection. Held-out labels entered only the final metric calculation.")
add_heading(doc, "4.2 Positive-unlabeled simulation", 2)
add_body(doc, "For each positive-label fraction, a seeded subset of deceptive training examples was revealed as P. All remaining positives and every negative became U. Ground-truth labels were removed from the training view; hidden labels were retained only by the experiment harness for evaluation. Consequently, the unlabeled pool remained a mixture rather than a pseudo-negative class.")
add_heading(doc, "4.3 nnPU objective", 2)
add_body(doc, "For score function f, logistic loss l, and positive class prior pi, the unbiased PU risk is pi E_P[l(f,+1)] + E_U[l(f,-1)] - pi E_P[l(f,-1)]. When the estimated negative-risk component becomes negative, the Kiryo non-negative correction reverses that component's gradient. The implementation returns positive, negative, corrected, and total risks and rejects invalid priors, shapes, empty batches, and non-finite logits.")
add_heading(doc, "4.4 PU-Mean Teacher detector", 2)
add_body(doc, "The combined model adds a sigmoid-ramped consistency loss to nnPU risk. A student receives stronger feature masking while an exponential-moving-average teacher receives a weaker view. Teacher probabilities contribute only when confidence exceeds 0.80. This differs from ordinary self-training: pseudo-label consistency supplements a valid PU objective and the unlabeled pool is never assigned a negative classification target.")
add_heading(doc, "4.5 Class-prior estimation", 2)
add_body(doc, "The notebook reported the true training prevalence only as a non-deployable oracle reference. A training-only Elkan-Noto-style estimate used out-of-fold probabilities for observed positive selection. Estimated priors increased from 0.063 at 1% revealed positives to 0.200 at 20%, while the oracle prevalence was 0.364. The gap demonstrates substantial prior-estimation bias in this quick-mode setup.")

add_heading(doc, "5. In-domain results", 1)
results = [
    ("Majority", "100%", "0.000", "0.389", "0.364", "0.500", "0.000"),
    ("Naive U-as-negative", "1%", "0.546", "0.390", "0.442", "0.588", "0.123"),
    ("nnPU", "1%", "0.544", "0.463", "0.457", "0.616", "0.134"),
    ("PU-Mean Teacher", "1%", "0.545", "0.412", "0.442", "0.607", "0.123"),
    ("Naive U-as-negative", "5%", "0.580", "0.552", "0.565", "0.690", "0.243"),
    ("nnPU", "5%", "0.602", "0.636", "0.550", "0.715", "0.314"),
    ("PU-Mean Teacher", "5%", "0.588", "0.580", "0.518", "0.682", "0.267"),
    ("Naive U-as-negative", "10%", "0.625", "0.629", "0.597", "0.741", "0.350"),
    ("nnPU", "10%", "0.628", "0.649", "0.603", "0.756", "0.360"),
    ("PU-Mean Teacher", "10%", "0.617", "0.643", "0.577", "0.733", "0.340"),
    ("Naive U-as-negative", "20%", "0.627", "0.668", "0.647", "0.769", "0.367"),
    ("nnPU", "20%", "0.653", "0.658", "0.632", "0.774", "0.408"),
    ("PU-Mean Teacher", "20%", "0.650", "0.658", "0.617", "0.768", "0.401"),
    ("Supervised TF-IDF upper bound", "100%", "0.712", "0.759", "0.775", "0.861", "0.528"),
]
add_table(doc, ["Method", "P revealed", "F1", "Macro F1", "PR-AUC", "ROC-AUC", "MCC"], results,
          [2900, 1100, 900, 1100, 1100, 1100, 1160], numeric_cols=(1, 2, 3, 4, 5, 6))
add_body(doc, "At matched label fractions, nnPU improved over the naive comparator most clearly at 5%, 10%, and 20% in F1, balanced ranking, or MCC, but not on every metric. At 10%, nnPU improved F1 by 0.003 and PR-AUC by 0.006. At 20%, it improved F1 by 0.026 and MCC by 0.041, while naive training retained a slightly higher PR-AUC. The combined consistency model did not surpass nnPU at any tested fraction on F1 or PR-AUC in this run.")
add_body(doc, "The 10% oracle-prior nnPU reference achieved F1 0.646 and PR-AUC 0.637, compared with F1 0.628 and PR-AUC 0.603 using the training-only estimate. Prior multipliers of 0.75 and 1.25 changed results materially. Positive-prior estimation is therefore not a peripheral detail; it is a central source of deployable-model uncertainty.")
if len(figures) > 3:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(figures[3], width=Inches(6.25))
    shape._inline.docPr.set("descr", "Line chart comparing PR-AUC across revealed-positive fractions for naive, nnPU, and PU-Mean Teacher models.")
    add_caption(doc, "Figure 3. Label-efficiency curves for the naive, nnPU, and PU-Mean Teacher models. Stored notebook output.")

add_heading(doc, "6. Cross-domain evaluation", 1)
lodo = [
    ("Job scams", "Supervised TF-IDF", "0.089", "0.173", "0.061", "0.620", "0.057"),
    ("Job scams", "PU-Mean Teacher 10%", "0.085", "0.173", "0.096", "0.639", "0.031"),
    ("Product reviews", "Supervised TF-IDF", "0.210", "0.416", "0.490", "0.478", "-0.034"),
    ("Product reviews", "PU-Mean Teacher 10%", "0.536", "0.495", "0.482", "0.474", "-0.004"),
    ("Phishing", "Supervised TF-IDF", "0.453", "0.495", "0.382", "0.505", "0.004"),
    ("Phishing", "PU-Mean Teacher 10%", "0.586", "0.458", "0.532", "0.646", "0.158"),
]
add_table(doc, ["Held-out domain", "Method", "F1", "Macro F1", "PR-AUC", "ROC-AUC", "MCC"], lodo,
          [1800, 2500, 900, 1100, 1000, 1000, 1060], numeric_cols=(2, 3, 4, 5, 6))
add_body(doc, "LODO results were substantially weaker than in-domain results. Job-scam F1 remained near 0.09, reflecting severe imbalance and distribution shift. Product-review ranking was near chance for both models despite the combined model's thresholded F1 of 0.536; its MCC remained approximately zero. Phishing was the strongest combined-model holdout, with ROC-AUC 0.646 and MCC 0.158, but false-positive rates remained operationally high. These results do not establish domain independence.")
add_callout(doc, "Interpretation", "Thresholded F1 can look favorable even when ranking and correlation metrics remain weak. Cross-domain conclusions must consider PR-AUC, ROC-AUC, MCC, prevalence, and false-positive burden together.", "FFF4E5")

add_heading(doc, "7. Error analysis and interpretability", 1)
add_body(doc, "The executed notebook stratified false positives and false negatives by domain, text length, and URL presence, and displayed a small set of truncated high-confidence errors. Classical TF-IDF coefficients highlighted URL artifacts, political names, product-review and job-posting templates, source-specific vocabulary, and formatting cues. These coefficients are associations rather than causal explanations. Attention weights were not used as definitive explanations.")
add_body(doc, "A bootstrap over test examples gave a 95% percentile interval of approximately 0.595 to 0.636 for the selected 10% PU-Mean Teacher F1. This interval reflects test-sample variation only; it does not capture seed, representation, hyperparameter, or domain uncertainty.")
if len(figures) > 4:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(figures[4], width=Inches(6.25))
    shape._inline.docPr.set("descr", "Three diagnostic panels showing precision-recall, calibration, and confusion matrix for the selected combined model.")
    add_caption(doc, "Figure 4. Precision-recall, calibration, and confusion-matrix diagnostics for the selected combined model. Stored notebook output.")

add_heading(doc, "8. Threats to validity and limitations", 1)
for item in (
    "Only one random seed was executed, so no training-variance estimate or valid paired significance test is available.",
    "Quick mode capped every domain/split at 500 stratified examples; its results must not be presented as full-dataset performance.",
    "PU labels were simulated by hiding known positives. Real labeling mechanisms may violate selected-completely-at-random assumptions.",
    "The Elkan-Noto-style prior estimate substantially underestimated the oracle prior, and the optimal prior may vary by domain.",
    "Near-duplicate and provenance artifacts can inflate in-domain performance and complicate attribution to deception semantics.",
    "Reduced LODO covered only job scams, product reviews, and phishing; all seven rotations remain necessary.",
    "Frozen sentence embeddings and small MLPs are computationally appropriate for quick mode but are not a substitute for a controlled GPU fine-tuning study.",
    "Deception detection is not factual verification and cannot establish speaker intent.",
): add_bullet(doc, item)

add_heading(doc, "9. Ethical and operational considerations", 1)
add_body(doc, "False positives may harm job applicants, political speakers, reviewers, or ordinary message senders. The observed cross-domain false-positive burdens make autonomous enforcement especially inappropriate. Any downstream system should operate as a human-reviewed alerting aid with calibrated domain-specific thresholds, audit logs, appeal mechanisms, privacy safeguards, subgroup evaluation, and drift monitoring.")
add_body(doc, "Dataset labels inherit the assumptions and provenance of constituent sources. A production claim requires source-by-source licensing and documentation review, external validation on post-collection data, and explicit analysis of whether protected or socially sensitive attributes are indirectly encoded by domain or source vocabulary.")

add_heading(doc, "10. Reproducibility and verification", 1)
checks = [
    ("Notebook execution", "14/14 code cells executed; zero stored error outputs"),
    ("Unit tests", "8 passed: finite loss, gradients, shapes, correction, invalid inputs, NaN safeguards"),
    ("Dataset pin", "aaaf94b336c563a14806bb4f3f58727bed9ed8d4"),
    ("Randomness", "Seed 42; deterministic PyTorch settings where practical"),
    ("Test-label isolation", "No fitting, prior estimation, early stopping, calibration, or threshold selection used test labels"),
    ("Artifact hygiene", "Data, caches, checkpoints, embeddings, outputs, and virtual environments are Git-ignored"),
]
add_table(doc, ["Control", "Evidence"], checks, [2500, 6860])

add_heading(doc, "11. Conclusions and recommended next experiments", 1)
add_body(doc, "Valid nnPU learning was competitive with and sometimes better than treating the unlabeled pool as negative, especially as more positives were revealed. However, the gains were modest and metric-dependent. Mean Teacher consistency did not add value consistently in the executed configuration. Strong domain separability, near-duplicate leakage risk, prior-estimation bias, and weak LODO results indicate that benchmark provenance remains a larger obstacle than in-domain classifier capacity.")
add_body(doc, "The next experiment should run three or more paired seeds across all four label fractions and all seven LODO rotations after cluster-level near-duplicate control. It should estimate domain-conditional priors, tune confidence and consistency schedules strictly within validation data, and compare frozen representations with compact-transformer fine-tuning on GPU. Paired permutation or Wilcoxon comparisons should be reported only after sufficient matched runs exist.")
add_callout(doc, "Bottom line", "The executed evidence supports PU learning as a principled alternative to pseudo-negative training, but it does not support a claim of domain-independent deception detection or a consistent benefit from Mean Teacher consistency.")

add_heading(doc, "References", 1)
references = [
    "DiFrauD dataset repository. Hugging Face: https://huggingface.co/datasets/difraud/difraud. Revision aaaf94b336c563a14806bb4f3f58727bed9ed8d4.",
    "Kiryo, R., Niu, G., du Plessis, M. C., and Sugiyama, M. (2017). Positive-Unlabeled Learning with Non-Negative Risk Estimator. Advances in Neural Information Processing Systems.",
    "Tarvainen, A., and Valpola, H. (2017). Mean Teachers Are Better Role Models: Weight-Averaged Consistency Targets Improve Semi-Supervised Deep Learning Results. Advances in Neural Information Processing Systems.",
    "Elkan, C., and Noto, K. (2008). Learning Classifiers from Only Positive and Unlabeled Data. Proceedings of KDD.",
]
for ref in references:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(.25); p.paragraph_format.first_line_indent = Inches(-.25); p.paragraph_format.space_after = Pt(5)
    font(p.add_run(ref), 9.5)

doc.core_properties.title = "Positive-Unlabeled and Mean Teacher Learning for Cross-Domain Deception Detection"
doc.core_properties.subject = "Technical report based on the executed DiFrauD quick-mode notebook"
doc.core_properties.author = "DiFrauD research project"
doc.core_properties.keywords = "DiFrauD, positive-unlabeled learning, nnPU, Mean Teacher, deception detection"
doc.save(OUTPUT)
print(OUTPUT)
